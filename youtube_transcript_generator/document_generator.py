#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""자막 문서 생성 모듈."""

import os
import re
from datetime import datetime
from typing import List, Dict, Any, Tuple, Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from youtube_transcript_generator.downloader import clean_filename
from youtube_transcript_generator.transcriber import format_time
from youtube_transcript_generator.translator import translate_text, translate_paragraphs, detect_language, TranslationError


def create_transcript_document(
    captions: List[Dict[str, Any]], 
    title: str, 
    output_dir: str,
    translate: bool = False,
    api_key: Optional[str] = None
) -> str:
    """자막을 문서 파일로 만들기."""
    # 타임스탬프 포함 자막과 전체 스크립트 문서를 각각 생성
    timestamp_file = create_timestamp_document(captions, title, output_dir)
    script_file = create_script_document(captions, title, output_dir, translate, api_key)
    
    # 타임스탬프 파일 경로 반환
    return timestamp_file


def create_timestamp_document(
    captions: List[Dict[str, Any]], 
    title: str, 
    output_dir: str
) -> str:
    """타임스탬프가 포함된 자막 문서 생성."""
    doc = Document()
    
    # 제목 추가
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 생성 시간 추가
    time_para = doc.add_paragraph(f"생성 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    time_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # 타임스탬프 포함 자막 섹션 추가
    doc.add_heading("타임스탬프 포함 자막", level=2)
    
    # 자막 내용 추가
    full_text = ""
    if captions:
        for caption in captions:
            time_str = format_time(caption['start'])
            text = caption['text']
            doc.add_paragraph(f"[{time_str}] {text}")
            full_text += text + " "
    else:
        doc.add_paragraph("자막이 없습니다.")
    
    # 자막이 거의 없는 경우 처리
    if not captions or len(full_text.strip()) < 30:
        warning_para = doc.add_paragraph("※ 이 동영상에는 충분한 자막이 없습니다.")
        warning_para.runs[0].bold = True
    
    # 파일 저장
    safe_title = clean_filename(title)
    output_file = os.path.join(output_dir, f"{safe_title}_타임스탬프.docx")
    doc.save(output_file)
    
    # 텍스트 파일로도 저장
    txt_output_file = os.path.join(output_dir, f"{safe_title}_타임스탬프.txt")
    with open(txt_output_file, 'w', encoding='utf-8') as f:
        f.write(f"{title}\n\n")
        f.write(f"생성 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        
        if not captions or len(full_text.strip()) < 30:
            f.write("※ 이 동영상에는 충분한 자막이 없습니다.\n\n")
        
        f.write("==== 타임스탬프 포함 자막 ====\n\n")
        if captions:
            for caption in captions:
                time_str = format_time(caption['start'])
                f.write(f"[{time_str}] {caption['text']}\n")
        else:
            f.write("자막이 없습니다.\n")
    
    return output_file


def create_script_document(
    captions: List[Dict[str, Any]], 
    title: str, 
    output_dir: str,
    translate: bool = False,
    api_key: Optional[str] = None
) -> str:
    """전체 스크립트 문서 생성 (문단 정리 포함)."""
    doc = Document()
    
    # 제목 추가
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 생성 시간 추가
    time_para = doc.add_paragraph(f"생성 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    time_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # 자막 텍스트 추출
    full_text = ""
    if captions:
        for caption in captions:
            full_text += caption['text'] + " "
    
    # 자막이 거의 없는 경우 처리
    if not captions or len(full_text.strip()) < 30:
        warning_para = doc.add_paragraph("※ 이 동영상에는 충분한 자막이 없습니다.")
        warning_para.runs[0].bold = True
        full_text = "※ 이 동영상에는 충분한 자막이 없습니다. 자동 생성된 자막이 제한적이거나 없는 경우입니다."
    
    # 전체 스크립트 정제
    refined_script = refine_script(full_text) if full_text.strip() else "자막이 없습니다."
    
    # # 문단별로 정리
    # paragraphs = split_into_paragraphs(full_text)
    
    # 언어 감지
    source_lang = detect_language(refined_script)
    
    # 번역이 필요한 경우 (언어가 한국어가 아닌 경우에만)
    translated_script = None
    translated_paragraphs = None
    
    if translate and source_lang != 'ko' and len(refined_script.strip()) > 30:
        try:
            print(f"감지된 언어: {source_lang}")
            print("전체 스크립트 번역 중...")
            translated_script = translate_text(refined_script, source_lang, 'ko', api_key)
            
            # if paragraphs:
            #     print("문단별 번역 중...")
            #     translated_paragraphs = translate_paragraphs(paragraphs, source_lang, 'ko', api_key)
        except TranslationError as e:
            print(f"번역 오류: {str(e)}")
            translated_script = None
            translated_paragraphs = None
    
    # 전체 스크립트 추가 (정제된 버전)
    doc.add_heading("전체 스크립트", level=2)
    full_script_para = doc.add_paragraph(refined_script)
    full_script_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # 번역된 스크립트가 있으면 추가
    if translated_script:
        doc.add_heading("번역된 전체 스크립트 (한국어)", level=2)
        trans_para = doc.add_paragraph(translated_script)
        trans_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # # 문단별로 정리된 버전 추가
    # doc.add_heading("문단별 정리 스크립트", level=2)
    
    # if paragraphs:
    #     for i, paragraph in enumerate(paragraphs):
    #         if paragraph.strip():
    #             para = doc.add_paragraph(paragraph)
    #             para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                
    #             # 번역된 문단이 있으면 추가
    #             if translated_paragraphs and i < len(translated_paragraphs):
    #                 trans_para = doc.add_paragraph()
    #                 trans_run = trans_para.add_run(f"[번역] {translated_paragraphs[i]}")
    #                 trans_run.italic = True
    #                 trans_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    
    #             # 구분선 추가 (마지막 문단 제외)
    #             if i < len(paragraphs) - 1:
    #                 doc.add_paragraph("---")
    # else:
    #     doc.add_paragraph("자막이 없거나 문단 분리가 불가능합니다.")
    
    # 파일 저장
    safe_title = clean_filename(title)
    output_file = os.path.join(output_dir, f"{safe_title}_전체스크립트.docx")
    doc.save(output_file)
    
    # 텍스트 파일로도 저장
    txt_output_file = os.path.join(output_dir, f"{safe_title}_전체스크립트.txt")
    with open(txt_output_file, 'w', encoding='utf-8') as f:
        f.write(f"{title}\n\n")
        f.write(f"생성 시간: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
        
        if not captions or len(full_text.strip()) < 30:
            f.write("※ 이 동영상에는 충분한 자막이 없습니다.\n\n")
        
        f.write("==== 전체 스크립트 ====\n\n")
        f.write(refined_script)
        
        if translated_script:
            f.write("\n\n==== 번역된 전체 스크립트 (한국어) ====\n\n")
            f.write(translated_script)
        
        # f.write("\n\n==== 문단별 정리 스크립트 ====\n\n")
        # if paragraphs:
        #     for i, paragraph in enumerate(paragraphs):
        #         if paragraph.strip():
        #             f.write(f"{paragraph}\n")
                    
        #             # 번역된 문단이 있으면 추가
        #             if translated_paragraphs and i < len(translated_paragraphs):
        #                 f.write(f"\n[번역] {translated_paragraphs[i]}\n")
                    
        #             f.write("\n---\n\n")
        # else:
        #     f.write("자막이 없거나 문단 분리가 불가능합니다.\n")
    
    return output_file


def refine_script(text: str) -> str:
    """전체 스크립트에서 중복되는 단어나 불필요한 표현을 정리."""
    # '[음악]' 같은 불필요한 표현 제거
    text = re.sub(r'\[음악\]', '', text)
    
    # 같은 문장이 반복되는 경우 한 번만 표시
    sentences = text.split('. ')
    unique_sentences = []
    for sentence in sentences:
        if sentence and sentence not in unique_sentences:
            unique_sentences.append(sentence)
    
    # 정리된 스크립트 반환
    refined_text = '. '.join(unique_sentences)
    
    # 마침표가 연속으로 나오는 경우 하나로 수정
    refined_text = re.sub(r'\.+', '.', refined_text)
    
    # 앞뒤 공백 제거
    refined_text = refined_text.strip()
    
    return refined_text


def split_into_paragraphs(text: str) -> List[str]:
    """전체 스크립트를 의미 단위로 문단 분리."""
    if not text or len(text.strip()) < 30:
        return []
    
    # 불필요한 표현 제거
    text = re.sub(r'\[음악\]', '', text)
    
    # 문장 분리 (마침표, 느낌표, 물음표 뒤 공백으로 구분)
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    # 문단 분리 (4-5개 문장을 하나의 문단으로)
    paragraphs = []
    current_paragraph = []
    for sentence in sentences:
        if sentence.strip():
            current_paragraph.append(sentence)
            
            # 문단 분리 조건: 문장 길이가 5 이상이거나 특정 키워드로 끝나는 경우
            if (len(current_paragraph) >= 5 or 
                any(keyword in sentence.lower() for keyword in ['그리고', '따라서', '그러나', '그래서', '결론적으로'])):
                paragraphs.append(' '.join(current_paragraph))
                current_paragraph = []
    
    # 남아있는 문장들도 문단으로 추가
    if current_paragraph:
        paragraphs.append(' '.join(current_paragraph))
    
    return paragraphs